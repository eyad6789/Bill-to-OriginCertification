#!/usr/bin/env python3
"""
Certificate of Origin Generator - Word Document Version
Uses .docx template for easy text replacement
Outputs both .docx and .pdf formats
"""

import os
import re
import json
import random
import subprocess
import shutil
from datetime import datetime, timedelta
from dataclasses import dataclass
from typing import List, Dict, Optional
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches
import pdfplumber

# For Gemini API
try:
    from google import genai
    USE_NEW_GENAI = True
except ImportError:
    try:
        import google.generativeai as genai
        USE_NEW_GENAI = False
    except ImportError:
        USE_NEW_GENAI = None


@dataclass
class BuyerInfo:
    name: str
    address: str
    mobile: str = ""
    tax_number: str = ""
    email: str = ""


@dataclass
class SellerInfo:
    name: str
    address: str


@dataclass
class ProductInfo:
    description: str
    hs_code: str
    quantity: str
    weight: str
    marks_numbers: str = "N/M"


@dataclass
class ShippingInfo:
    port_of_loading: str
    port_of_discharge: str
    destination_country: str


@dataclass
class InvoiceInfo:
    invoice_number: str
    invoice_date: str


@dataclass
class CertificateData:
    buyer: BuyerInfo
    seller: SellerInfo
    product: ProductInfo
    shipping: ShippingInfo
    invoice: InvoiceInfo
    certificate_number: str = ""
    serial_number: str = ""
    declaration_date: str = ""


class WordCertificateGenerator:
    """
    Generate Certificate of Origin using Word document template
    Much cleaner than PDF manipulation - just find and replace text!
    """
    
    # Define placeholders to search for and replace
    # These are the actual values from the template that will be replaced
    REPLACEMENTS = {
        # Exporter info
        'exporter_name': 'Yiwu Kabul Daily Necessities Factory',
        'exporter_address': 'ShowRoom 602, the 6th Floor, No.520, Dafuzhai Village, Houzhai Sub-dist, Yiwu City, Jinhua City, Zhejiang Province',
        
        # Certificate numbers
        'serial_no': 'CCPIT3512500229074',
        'cert_no': '25C351120866/00014',
        
        # Consignee info
        'consignee_name': 'ASHURBANIPAL COMPANY FOR GENERAL TRADE IN ELECTRICAL APPLIANCES AND HOME AND OFFICE FURNITURE, TRADE AND SUPPLY OF',
        'consignee_name_2': 'KITCHENS, FURNITURE AND DECORATION, PROCESSING AND MARKETING OF SINGLE-USE HOUSEHOLD AND FOOD SUPPLIES.',
        'consignee_address': 'ADDRESS : IRAQ - BAGHDAD / AL-QADISIYAH DISTRICT - MAHALLA / 606 - ALLEY',
        'consignee_address_2': '/ 8 - BUILDING NO. / 74 TABARAK CENTER BUILDING FLOOR NO. 5, OFFICE NO. 9',
        
        # Transport
        'transport_route': 'FROM NINGBO CHINA TO UMM QASR IRAQ BY SEA',
        
        # Destination
        'destination': 'IRAQ',
        
        # Marks
        'marks': 'N/M',
        
        # Product description
        'product_desc_1': 'SIX HUNDRED FORTY',
        'product_desc_2': '(640) CTNS OF GLASS ELECTRIC KETTLE',
        
        # Contact info
        'mobile': '009647901860410 - 009647844702070',
        'tax_number': '902191163',
        'email': '***',  # Email placeholder in template
        
        # HS Code
        'hs_code': '851671.00',
        
        # Weight
        'weight': '7,910 KGS G.W.',
        
        # Invoice
        'invoice_no_1': 'YKDNASH71374',
        'invoice_no_2': '93',
        'invoice_date': 'OCT.09,2025',
        
        # Declaration dates
        'declaration_date': 'NOV.25,2025',
        
        # Labels to remove (will be replaced with empty string - include newline to remove whitespace)
        'signature_label_1': '\nPlace and date,signature and stamp of certifying authority',
        'signature_label_2': '\nPlace and date,signature and stamp of authorized signatory',
    }
    
    def __init__(self, template_path: str):
        """
        Initialize with Word template
        
        Args:
            template_path: Path to the .docx template file
        """
        self.template_path = template_path
    
    def generate_declaration_date(self, invoice_date_str: str) -> str:
        """Generate declaration date 10-30 days after invoice date"""
        try:
            invoice_date = datetime.strptime(invoice_date_str, "%b.%d,%Y")
        except ValueError:
            try:
                invoice_date = datetime.strptime(invoice_date_str, "%d-%b-%Y")
            except ValueError:
                try:
                    invoice_date = datetime.strptime(invoice_date_str, "%Y-%m-%d")
                except ValueError:
                    invoice_date = datetime.now()
        
        random_days = random.randint(10, 30)
        declaration_date = invoice_date + timedelta(days=random_days)
        return declaration_date.strftime("%b.%d,%Y").upper()
    
    def generate_certificate_number(self) -> tuple:
        """Generate certificate and serial numbers"""
        serial = f"CCPIT351250{random.randint(100000, 999999)}"
        cert = f"25C35112{random.randint(1000, 9999)}/000{random.randint(10, 99)}"
        return serial, cert
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """
        Replace text in a paragraph while preserving formatting
        Returns True if replacement was made
        """
        # Get full paragraph text
        full_text = ''.join(run.text for run in paragraph.runs)
        
        if old_text in full_text:
            # Find which runs contain the text
            new_full_text = full_text.replace(old_text, new_text)
            
            # Clear all runs and set new text in first run
            if paragraph.runs:
                # Preserve formatting from first run
                first_run = paragraph.runs[0]
                
                # Clear all runs
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = new_full_text
                    else:
                        run.text = ''
                
                return True
        
        return False
    
    def _replace_in_document(self, doc: Document, old_text: str, new_text: str):
        """Replace text throughout the document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, old_text, new_text)
    
    def create_certificate(self, data: CertificateData, output_path: str, output_pdf: bool = True):
        """
        Create certificate by replacing text in template
        
        Args:
            data: CertificateData with all information
            output_path: Output file path (.docx)
            output_pdf: Also create PDF version
        """
        # Load template
        doc = Document(self.template_path)
        
        # Build replacement mapping
        replacements = {
            # Exporter
            self.REPLACEMENTS['exporter_name']: data.seller.name,
            self.REPLACEMENTS['exporter_address']: data.seller.address,
            
            # Certificate numbers
            self.REPLACEMENTS['serial_no']: data.serial_number,
            self.REPLACEMENTS['cert_no']: data.certificate_number,
            
            # Consignee - need to handle multi-line carefully
            self.REPLACEMENTS['consignee_name']: self._get_consignee_name_part1(data.buyer.name),
            self.REPLACEMENTS['consignee_name_2']: self._get_consignee_name_part2(data.buyer.name),
            self.REPLACEMENTS['consignee_address']: f"ADDRESS : {self._get_address_part1(data.buyer.address)}",
            self.REPLACEMENTS['consignee_address_2']: self._get_address_part2(data.buyer.address),
            
            # Transport
            self.REPLACEMENTS['transport_route']: f"FROM {data.shipping.port_of_loading} TO {data.shipping.port_of_discharge} BY SEA",
            
            # Destination (appears multiple times)
            'IRAQ': data.shipping.destination_country,
            
            # Product
            self.REPLACEMENTS['product_desc_1']: self._get_product_part1(data.product.description),
            self.REPLACEMENTS['product_desc_2']: self._get_product_part2(data.product.description),
            
            # Contact info
            self.REPLACEMENTS['mobile']: data.buyer.mobile if data.buyer.mobile else "N/A",
            self.REPLACEMENTS['tax_number']: data.buyer.tax_number if data.buyer.tax_number else "N/A",
            self.REPLACEMENTS['email']: f"EMAIL : {data.buyer.email}" if data.buyer.email else "",
            
            # HS Code
            self.REPLACEMENTS['hs_code']: data.product.hs_code,
            
            # Weight
            self.REPLACEMENTS['weight']: data.product.weight,
            
            # Invoice
            self.REPLACEMENTS['invoice_no_1']: data.invoice.invoice_number[:12] if len(data.invoice.invoice_number) > 12 else data.invoice.invoice_number,
            self.REPLACEMENTS['invoice_no_2']: data.invoice.invoice_number[12:] if len(data.invoice.invoice_number) > 12 else "",
            self.REPLACEMENTS['invoice_date']: data.invoice.invoice_date,
            
            # Declaration date (appears twice)
            self.REPLACEMENTS['declaration_date']: data.declaration_date,
        }
        
        # Perform replacements
        for old_text, new_text in replacements.items():
            if old_text and new_text is not None:
                self._replace_in_document(doc, old_text, new_text)
        
        # Remove signature/stamp labels (replace with empty string)
        self._replace_in_document(doc, self.REPLACEMENTS['signature_label_1'], '')
        self._replace_in_document(doc, self.REPLACEMENTS['signature_label_2'], '')
        
        # Handle marks separately (it's just "N/M" which appears in headers too)
        # Only replace in the specific table cell
        
        # Save Word document
        doc.save(output_path)
        print(f"Word document generated: {output_path}")
        
        # Convert to PDF if requested
        if output_pdf:
            pdf_path = output_path.rsplit('.', 1)[0] + '.pdf'
            self._convert_to_pdf(output_path, pdf_path)
            return output_path, pdf_path
        
        return output_path, None
    
    def _get_consignee_name_part1(self, full_name: str) -> str:
        """Get first part of consignee name (fits first line)"""
        words = full_name.split()
        line = ""
        for word in words:
            test = line + (" " if line else "") + word
            if len(test) > 75:
                break
            line = test
        return line
    
    def _get_consignee_name_part2(self, full_name: str) -> str:
        """Get second part of consignee name"""
        part1 = self._get_consignee_name_part1(full_name)
        remaining = full_name[len(part1):].strip()
        return remaining if remaining else ""
    
    def _get_address_part1(self, address: str) -> str:
        """Get first part of address"""
        if len(address) <= 60:
            return address
        # Find a good break point
        words = address.split()
        line = ""
        for word in words:
            test = line + (" " if line else "") + word
            if len(test) > 60:
                break
            line = test
        return line
    
    def _get_address_part2(self, address: str) -> str:
        """Get second part of address"""
        part1 = self._get_address_part1(address)
        remaining = address[len(part1):].strip()
        if remaining.startswith('-'):
            remaining = remaining[1:].strip()
        return remaining if remaining else ""
    
    def _get_product_part1(self, description: str) -> str:
        """Get first part of product description"""
        # Usually quantity in words
        words = description.split()
        line = ""
        for word in words:
            test = line + (" " if line else "") + word
            if len(test) > 20 or '(' in word:
                break
            line = test
        return line
    
    def _get_product_part2(self, description: str) -> str:
        """Get second part of product description"""
        part1 = self._get_product_part1(description)
        return description[len(part1):].strip()
    
    def _convert_to_pdf(self, docx_path: str, pdf_path: str):
        """Convert Word document to PDF - uses docx2pdf on Windows, LibreOffice on Linux"""
        import platform
        
        # Try docx2pdf first (Windows with Microsoft Word installed)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            print(f"PDF generated: {pdf_path}")
            return True
        except ImportError:
            pass
        except Exception as e:
            print(f"docx2pdf failed: {e}")
        
        # Fallback to LibreOffice (Linux/Mac or if Word not available)
        try:
            output_dir = os.path.dirname(pdf_path) or '.'
            subprocess.run([
                'soffice', '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ], check=True, capture_output=True)
            print(f"PDF generated: {pdf_path}")
            return True
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            print(f"Warning: Could not convert to PDF: {e}")
            if platform.system() == 'Windows':
                print("Make sure Microsoft Word is installed for PDF conversion")
            else:
                print("Install LibreOffice for PDF conversion: apt install libreoffice")
            return False



class GeminiExtractor:
    """Extract data from Bill of Lading using Google Gemini API with automatic model fallback"""
    
    # Models to try in order of preference
    MODELS = [
        'gemini-2.5-flash',
        'gemini-2.5-flash-lite',
        'gemini-3-flash',
        'gemini-1.5-flash',
        'gemini-1.5-flash-latest',
        'gemini-pro',
        'gemma-3-27b',
        'gemma-3-12b',
        'gemma-3-4b',
        'gemma-3-2b',
        'gemma-3-1b',
    ]
    
    def __init__(self, api_key: str = None):
        """Initialize with API key - loads from .env if not provided"""
        if api_key is None:
            from dotenv import load_dotenv
            load_dotenv()
            api_key = os.environ.get('GEMINI_API_KEY')
        
        if not api_key or api_key == 'your_api_key_here':
            raise ValueError("GEMINI_API_KEY not set. Please set it in .env file or pass directly.")
        
        self.api_key = api_key
        
        if USE_NEW_GENAI is True:
            self.client = genai.Client(api_key=api_key)
        elif USE_NEW_GENAI is False:
            genai.configure(api_key=api_key)
        else:
            raise ImportError("Google Generative AI package not installed")
    
    def _call_with_fallback(self, prompt, image=None):
        """Try each model until one works"""
        last_error = None
        
        for model_name in self.MODELS:
            try:
                print(f"Trying model: {model_name}")
                
                if USE_NEW_GENAI:
                    if image:
                        response = self.client.models.generate_content(
                            model=model_name,
                            contents=[prompt, image]
                        )
                    else:
                        response = self.client.models.generate_content(
                            model=model_name,
                            contents=prompt
                        )
                else:
                    model = genai.GenerativeModel(model_name)
                    if image:
                        response = model.generate_content([prompt, image])
                    else:
                        response = model.generate_content(prompt)
                
                # If we got here, the model worked
                print(f"Success with model: {model_name}")
                return response
                
            except Exception as e:
                last_error = e
                print(f"Model {model_name} failed: {str(e)[:100]}")
                continue
        
        # All models failed
        raise Exception(f"All models failed. Last error: {last_error}")
    
    def extract_from_bill(self, pdf_path: str) -> dict:
        text_content = self._extract_pdf_text(pdf_path)
        
        if not text_content.strip():
            return self._extract_from_pdf_image(pdf_path)
        
        return self._extract_from_text(text_content)
    
    def _extract_from_text(self, text_content: str) -> dict:
        prompt = self._get_extraction_prompt(f"Bill of Lading Content:\n{text_content}")
        response = self._call_with_fallback(prompt)
        return self._parse_response(response)
    
    def _extract_from_pdf_image(self, pdf_path: str) -> dict:
        """Extract from PDF by converting to image using PyMuPDF (pure Python, no external deps)"""
        
        # Try PyMuPDF first (pure Python, works everywhere)
        try:
            print(f"Trying PyMuPDF conversion for: {pdf_path}")
            import fitz  # PyMuPDF
            import io
            import PIL.Image
            
            doc = fitz.open(pdf_path)
            if len(doc) > 0:
                page = doc[0]  # First page
                # Render at higher resolution for better OCR
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img_data = pix.tobytes("png")
                image = PIL.Image.open(io.BytesIO(img_data))
                
                print(f"PyMuPDF: Converted PDF to {image.size[0]}x{image.size[1]} image")
                
                prompt = self._get_extraction_prompt("Analyze this Bill of Lading image and extract all shipping information.")
                response = self._call_with_fallback(prompt, image)
                result = self._parse_response(response)
                if result:
                    print("Success with PyMuPDF image extraction!")
                    return result
            doc.close()
        except Exception as e:
            print(f"PyMuPDF extraction failed: {str(e)[:100]}")
        
        # Fallback: Try Gemini native PDF upload
        try:
            print("Fallback: Trying Gemini native PDF upload...")
            if USE_NEW_GENAI:
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()
                
                from google.genai import types
                prompt = self._get_extraction_prompt("Extract all information from this Bill of Lading PDF document.")
                
                for model_name in self.MODELS:
                    try:
                        print(f"Trying model with PDF: {model_name}")
                        response = self.client.models.generate_content(
                            model=model_name,
                            contents=[
                                types.Content(
                                    parts=[
                                        types.Part.from_bytes(data=pdf_content, mime_type='application/pdf'),
                                        types.Part.from_text(text=prompt)
                                    ]
                                )
                            ]
                        )
                        result = self._parse_response(response)
                        if result:
                            print(f"Success with PDF upload using model: {model_name}")
                            return result
                    except Exception as e:
                        print(f"PDF upload failed with {model_name}: {str(e)[:100]}")
                        continue
        except Exception as e:
            print(f"Gemini native PDF upload failed: {str(e)[:100]}")
        
        return {}
    
    def _get_extraction_prompt(self, context: str) -> str:
        return f"""
Extract shipping document information as JSON:

{{
    "buyer": {{"name": "", "address": "", "mobile": "", "tax_number": "", "email": ""}},
    "seller": {{"name": "", "address": ""}},
    "product": {{"description": "", "hs_code": "", "quantity": "", "weight": "", "marks_numbers": "N/M"}},
    "shipping": {{"port_of_loading": "", "port_of_discharge": "", "destination_country": ""}},
    "invoice": {{"invoice_number": "", "invoice_date": ""}}
}}

IMPORTANT DATA SOURCE RULES:
- From BILL OF LADING: buyer name, buyer address, seller name, seller address, product description, HS code, quantity, weight, shipping ports, destination country
- From INVOICE ONLY: invoice number and invoice date
- Date format: MMM.DD,YYYY (e.g., OCT.09,2025)

Return ONLY valid JSON, no explanation.

{context}
"""
    
    def _parse_response(self, response) -> dict:
        try:
            text = response.text.strip()
            if text.startswith("```"):
                text = re.sub(r'^```json?\n?', '', text)
                text = re.sub(r'\n?```$', '', text)
            return json.loads(text)
        except:
            return {}
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            pass
        return text


def process_bill_of_lading(bill_pdf_path: str, api_key: str, template_path: str, output_path: str = None):
    """Process Bill of Lading and generate Certificate"""
    
    extractor = GeminiExtractor(api_key)
    
    print(f"Extracting from: {bill_pdf_path}")
    data = extractor.extract_from_bill(bill_pdf_path)
    
    if not data:
        raise ValueError("Failed to extract data")
    
    print(json.dumps(data, indent=2))
    
    buyer = BuyerInfo(**data.get('buyer', {}))
    seller = SellerInfo(**data.get('seller', {}))
    product = ProductInfo(**data.get('product', {}))
    shipping = ShippingInfo(**data.get('shipping', {}))
    invoice = InvoiceInfo(**data.get('invoice', {}))
    
    generator = WordCertificateGenerator(template_path)
    serial_no, cert_no = generator.generate_certificate_number()
    declaration_date = generator.generate_declaration_date(invoice.invoice_date)
    
    cert_data = CertificateData(
        buyer=buyer,
        seller=seller,
        product=product,
        shipping=shipping,
        invoice=invoice,
        serial_number=serial_no,
        certificate_number=cert_no,
        declaration_date=declaration_date
    )
    
    if not output_path:
        output_path = f"certificate_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    return generator.create_certificate(cert_data, output_path)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Certificate of Origin')
    parser.add_argument('bill_pdf', help='Bill of Lading PDF')
    parser.add_argument('--api-key', required=True, help='Gemini API key')
    parser.add_argument('--template', '-t', required=True, help='Word template (.docx)')
    parser.add_argument('--output', '-o', help='Output path')
    parser.add_argument('--no-pdf', action='store_true', help='Skip PDF generation')
    
    args = parser.parse_args()
    
    docx_path, pdf_path = process_bill_of_lading(
        args.bill_pdf,
        args.api_key,
        args.template,
        args.output
    )
    
    print(f"\nGenerated: {docx_path}")
    if pdf_path:
        print(f"Generated: {pdf_path}")


if __name__ == "__main__":
    main()
